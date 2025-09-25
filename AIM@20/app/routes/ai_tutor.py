from flask import Blueprint, render_template, request, jsonify, flash, redirect, url_for
from flask_login import login_required, current_user
from app import db
from app.models import AIChat, Quote, Task, Goal, Quiz, QuizQuestion, QuizAttempt, QuizAnswer
import json
import random
import re
from datetime import datetime, timedelta
import os
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from dotenv import load_dotenv
from flask import session
from werkzeug.utils import secure_filename
import logging

ai_tutor_bp = Blueprint('ai_tutor', __name__, url_prefix='/ai-tutor')

class GeminiTutor:
    """AI Study Assistant following working FarmingAssistant pattern"""

    def __init__(self):
        # Load environment variables
        load_dotenv()

        # Check for API key
        api_key = os.getenv('GEMINI_API_KEY', '').strip()

        if not api_key or api_key == '' or api_key == 'AIzaSyB9Q8w8k8Q8w8k8Q8w8k8Q8w8k8Q8w8k8Q8w8k' or len(api_key) < 20:
            print("WARNING: Gemini API key not configured properly. Using fallback content only.")
            self.api_available = False
            self.llm = None
            self.conversation = None
        else:
            try:
                # Configure Gemini API (following working Streamlit pattern)
                genai.configure(api_key=api_key)

                # Set generation configuration (matching working example)
                generation_config = {
                    "temperature": 0.8,
                    "top_k": 40,
                    "top_p": 0.8,
                    "max_output_tokens": 2048,
                }

                # Initialize Gemini model directly (following working pattern)
                self.model = genai.GenerativeModel(
                    model_name="gemini-2.5-flash-lite",
                    generation_config=generation_config
                )

                # Test the API with a simple request
                try:
                    test_response = self.model.generate_content("Hello, Gemini API test")
                    print("SUCCESS: Gemini API test successful")
                except Exception as test_e:
                    print(f"WARNING: Gemini API test failed: {str(test_e)}")
                    # Continue anyway, might work for actual requests

                # Initialize LangChain with Gemini for conversational features
                self.llm = ChatGoogleGenerativeAI(
                    model="gemini-2.5-flash-lite",
                    temperature=0.8,
                    top_k=40,
                    top_p=0.8,
                    google_api_key=api_key,
                    request_timeout=30,
                    max_retries=2
                )

                # Create conversation memory
                self.memory = ConversationBufferMemory()

                # Initialize conversation chain
                self.conversation = ConversationChain(
                    llm=self.llm,
                    memory=self.memory,
                    verbose=False
                )

                self.api_available = True
                print("SUCCESS: Gemini API initialized successfully with proper configuration")

            except Exception as e:
                print(f"ERROR: Failed to initialize Gemini API: {str(e)}")
                self.api_available = False
                self.model = None
                self.llm = None
                self.conversation = None

        # Study tips for the study_tip route
        self.study_tips = [
            "Break your study sessions into 25-minute focused intervals with 5-minute breaks (Pomodoro technique).",
            "Use active recall: Test yourself on the material rather than just re-reading.",
            "Space out your study sessions over time rather than cramming (spaced repetition).",
            "Teach the material to someone else to reinforce your understanding.",
            "Get enough sleep - your brain consolidates memories during sleep.",
            "Stay hydrated and eat brain-boosting foods like nuts, berries, and dark chocolate.",
            "Exercise regularly to improve cognitive function and reduce stress.",
            "Create mind maps to visualize connections between concepts."
        ]

    def get_study_guidance(self, topic=None, question=None):
        """Get comprehensive study guidance"""

        if question:
            # Handle specific user questions
            return self._answer_question(question)
        elif topic:
            # Provide guidance for specific study topics
            return self._get_topic_guidance(topic)
        else:
            # Provide complete study guide
            return self._get_complete_guide()

    def _get_complete_guide(self):
        """Get complete study guide"""

        # Check if API is available
        if not self.api_available or not self.model:
            return {
                'success': False,
                'error': 'Gemini API not configured. Please set GEMINI_API_KEY in .env file.',
                'fallback_content': self._get_fallback_guide()
            }

        prompt = f"""
        You are an expert study coach and academic advisor.
        Provide a comprehensive, step-by-step guide for effective studying and academic success.

        Structure your response with clear sections and practical tips. Include:
        1. Study environment setup
        2. Time management and scheduling
        3. Note-taking techniques
        4. Active learning strategies
        5. Memory and retention methods
        6. Exam preparation
        7. Stress management
        8. Common study mistakes to avoid

        Use simple language and include practical, actionable advice.
        Keep the response under 2000 words.
        """

        try:
            # Use direct Gemini API (following working Streamlit pattern)
            response = self.model.generate_content(prompt)
            return {
                'success': True,
                'content': response.text,
                'type': 'complete_guide'
            }
        except Exception as e:
            error_msg = str(e).lower()
            if 'timeout' in error_msg or 'connection' in error_msg or 'network' in error_msg:
                user_friendly_error = 'Connection timeout. Check your internet connection and try again.'
            elif 'quota' in error_msg or 'rate limit' in error_msg:
                user_friendly_error = 'API quota exceeded. Please try again later.'
            elif 'invalid' in error_msg or 'unauthorized' in error_msg:
                user_friendly_error = 'API key issue. Please check your GEMINI_API_KEY configuration.'
            else:
                user_friendly_error = f'AI service temporarily unavailable: {str(e)}'

            return {
                'success': False,
                'error': user_friendly_error,
                'fallback_content': self._get_fallback_guide()
            }

    def _get_topic_guidance(self, topic):
        """Get guidance for specific study topics"""

        # Check if API is available
        if not self.api_available or not self.model:
            return {
                'success': False,
                'error': 'Gemini API not configured. Please set GEMINI_API_KEY in .env file.',
                'topic': topic,
                'fallback_content': self._get_topic_fallback(topic)
            }

        topic_prompts = {
            'time_management': "Guide on effective time management for students",
            'note_taking': "Best practices for note-taking and organization",
            'exam_prep': "Strategies for exam preparation and test-taking",
            'motivation': "How to stay motivated and overcome procrastination",
            'memory': "Memory techniques and retention strategies",
            'math': "Study strategies for mathematics",
            'science': "Effective learning methods for science subjects",
            'languages': "Techniques for learning foreign languages",
            'writing': "Skills for academic writing and research"
        }

        prompt = f"""
        You are an academic coach specializing in study skills. Provide detailed, practical guidance for studying {topic}.
        Make it specific, actionable, and suitable for students.

        Focus on:
        - Best practices and techniques
        - Common mistakes to avoid
        - Resources and tools to use
        - Step-by-step approaches
        - Tips for different learning styles

        Keep the response under 1500 words.
        Topic: {topic}
        Guidance needed: {topic_prompts.get(topic, f'General study guidance for {topic}')}
        """

        try:
            # Use direct Gemini API (following working Streamlit pattern)
            response = self.model.generate_content(prompt)
            return {
                'success': True,
                'content': response.text,
                'type': 'topic_guide',
                'topic': topic
            }
        except Exception as e:
            error_msg = str(e).lower()
            if 'timeout' in error_msg or 'connection' in error_msg or 'network' in error_msg:
                user_friendly_error = 'Connection timeout. Check your internet connection and try again.'
            elif 'quota' in error_msg or 'rate limit' in error_msg:
                user_friendly_error = 'API quota exceeded. Please try again later.'
            elif 'invalid' in error_msg or 'unauthorized' in error_msg:
                user_friendly_error = 'API key issue. Please check your GEMINI_API_KEY configuration.'
            else:
                user_friendly_error = f'AI service temporarily unavailable: {str(e)}'

            return {
                'success': False,
                'error': user_friendly_error,
                'topic': topic,
                'fallback_content': self._get_topic_fallback(topic)
            }

    def _answer_question(self, question, language='en'):
        """Answer specific study questions using conversational AI"""

        # Check if API is available
        if not self.api_available or not self.model:
            return {
                'success': False,
                'error': 'Gemini API not configured. Please set GEMINI_API_KEY in .env file.',
                'question': question,
                'fallback_content': self._get_question_fallback(question, language)
            }

        # Language instructions
        language_instructions = {
            'en': 'Respond in English.',
            'es': 'Responde en español.',
            'fr': 'Réponds en français.',
            'de': 'Antworte auf Deutsch.',
            'pt': 'Responda em português.',
            'zh': '用中文回答。',
            'ja': '日本語で答えてください。',
            'ar': 'أجب باللغة العربية.',
            'hi': 'हिंदी में उत्तर दें।'
        }

        lang_instruction = language_instructions.get(language, 'Respond in English.')

        # Detect if this is a simple greeting or casual message
        question_lower = question.lower().strip()
        is_greeting = any(greeting in question_lower for greeting in [
            'hi', 'hello', 'hey', 'good morning', 'good afternoon', 'good evening',
            'how are you', 'what\'s up', 'howdy', 'greetings', 'salutations',
            'hola', 'bonjour', 'guten tag', 'buenos dias', 'ciao', 'aloha'
        ]) and len(question.split()) <= 5

        if is_greeting:
            # Conversational response for greetings
            context_prompt = f"""
            You are a friendly academic advisor and study coach helping students.

            The student said: "{question}"

            {lang_instruction}

            Guidelines for greetings:
            - Be warm and welcoming
            - Keep it conversational and brief (under 100 words)
            - Show enthusiasm about helping with studies
            - Don't give detailed study advice unless specifically asked
            - End by inviting them to ask study-related questions

            Respond naturally as if in a conversation.
            """
        else:
            # Detailed response for actual questions
            context_prompt = f"""
            You are a knowledgeable academic advisor and study coach helping students.
            Answer the student's question: "{question}"

            {lang_instruction}

            Guidelines:
            - Be practical and specific
            - Use simple, clear language appropriate for the selected language
            - Include evidence-based study techniques
            - Suggest affordable, accessible solutions
            - Include preventive measures for common study issues
            - Mention when to seek additional academic support

            Keep the response under 1000 words.
            Provide actionable advice based on proven study methods.
            """

        try:
            # Use direct Gemini API (following working Streamlit pattern)
            response = self.model.generate_content(context_prompt)
            return {
                'success': True,
                'content': response.text,
                'type': 'question_answer',
                'question': question
            }
        except Exception as e:
            error_msg = str(e).lower()
            if 'timeout' in error_msg or 'connection' in error_msg or 'network' in error_msg:
                user_friendly_error = 'Connection timeout. Check your internet connection and try again.'
            elif 'quota' in error_msg or 'rate limit' in error_msg:
                user_friendly_error = 'API quota exceeded. Please try again later.'
            elif 'invalid' in error_msg or 'unauthorized' in error_msg:
                user_friendly_error = 'API key issue. Please check your GEMINI_API_KEY configuration.'
            else:
                user_friendly_error = f'AI service temporarily unavailable: {str(e)}'

            return {
                'success': False,
                'error': user_friendly_error,
                'question': question,
                'fallback_content': self._get_question_fallback(question)
            }

    def get_response(self, user_message, message_type='general', context='', language='en'):
        """Get AI response - simplified for chat interface"""
        # Check if user has uploaded documents for context
        document_context = self._get_document_context()

        if document_context:
            # Include document context in the question
            enhanced_message = f"Based on the uploaded document(s), {user_message}\n\nDocument context: {document_context}"
            return self._answer_question(enhanced_message, language)
        else:
            return self._answer_question(user_message, language)

    def _get_document_context(self):
        """Get context from uploaded documents"""
        try:
            from flask import session
            documents = session.get('uploaded_documents', [])
            if documents:
                # Combine content from all uploaded documents
                contexts = []
                for doc in documents:  # Use all uploaded documents for context
                    if 'content' in doc:
                        contexts.append(f"From {doc['filename']}: {doc['content'][:3000]}...")

                return '\n\n'.join(contexts) if contexts else None
        except:
            pass
        return None

    def _get_fallback_guide(self):
        """Fallback study guide if AI fails"""
        return """
        # Complete Study Success Guide

        ## 1. Study Environment Setup
        - Find a quiet, well-lit space dedicated to studying
        - Keep your desk organized and free from distractions
        - Ensure good ventilation and comfortable temperature
        - Have all necessary materials within reach

        ## 2. Time Management
        - Use the Pomodoro technique: 25 minutes study + 5 minutes break
        - Create a weekly schedule with specific study blocks
        - Prioritize difficult subjects during peak energy times
        - Include buffer time for unexpected events

        ## 3. Note-Taking Techniques
        - Use the Cornell method for structured notes
        - Highlight key concepts and create summaries
        - Draw diagrams and mind maps for visual learners
        - Review and rewrite notes within 24 hours

        ## 4. Active Learning Strategies
        - Practice active recall: test yourself on material
        - Use spaced repetition for better retention
        - Teach concepts to others to reinforce understanding
        - Apply knowledge through practice problems

        ## 5. Memory & Retention
        - Get adequate sleep (7-9 hours) for memory consolidation
        - Stay hydrated and eat brain-boosting foods
        - Exercise regularly to improve cognitive function
        - Use mnemonic devices for difficult information

        ## 6. Exam Preparation
        - Start early - avoid last-minute cramming
        - Focus on understanding concepts, not memorization
        - Practice with past exams and sample questions
        - Get adequate sleep before the exam

        ## 7. Stress Management
        - Take regular breaks to prevent burnout
        - Practice deep breathing and meditation
        - Maintain a healthy work-life balance
        - Seek support when feeling overwhelmed

        ## Common Mistakes to Avoid
        - Multitasking during study sessions
        - Highlighting everything instead of key points
        - Studying in marathon sessions without breaks
        - Neglecting sleep and self-care
        """

    def _get_topic_fallback(self, topic):
        """Get fallback content for specific study topics"""
        fallbacks = {
            'time_management': """
            ## Time Management for Students
            - Use a planner or digital calendar for scheduling
            - Prioritize tasks using Eisenhower Matrix
            - Break large assignments into smaller tasks
            - Set realistic deadlines and stick to them
            - Avoid procrastination with the 5-minute rule
            """,
            'note_taking': """
            ## Effective Note-Taking
            - Use abbreviations and symbols to write faster
            - Leave margins for additional notes and questions
            - Organize notes by topic and date
            - Review and organize notes regularly
            - Use color coding for different types of information
            """,
            'exam_prep': """
            ## Exam Preparation Strategies
            - Create a study timeline starting 2-3 weeks before
            - Focus on understanding rather than memorization
            - Practice with past exams and sample questions
            - Get adequate sleep before the exam
            - Arrive early and stay calm during the test
            """
        }
        return fallbacks.get(topic, f"General study guidance for {topic} not available offline.")

    def _get_question_fallback(self, question):
        """Get fallback answer for questions"""
        question_lower = question.lower()

        if 'time' in question_lower or 'schedule' in question_lower:
            return """
            ## Time Management Advice:
            - Create a daily schedule with fixed study blocks
            - Use the Pomodoro technique for focused work
            - Prioritize tasks by importance and deadlines
            - Include breaks and self-care in your schedule
            - Track your time to identify productivity patterns
            """
        elif 'memory' in question_lower or 'remember' in question_lower:
            return """
            ## Memory Improvement Tips:
            - Use active recall: test yourself regularly
            - Practice spaced repetition for long-term retention
            - Create associations and mnemonic devices
            - Get enough sleep for memory consolidation
            - Stay physically active and eat brain-healthy foods
            """
        elif 'motivation' in question_lower or 'procrastination' in question_lower:
            return """
            ## Staying Motivated:
            - Set small, achievable goals and celebrate wins
            - Find your 'why' - connect studies to future goals
            - Create accountability through study partners
            - Reward yourself for completing study sessions
            - Visualize success and maintain a positive mindset
            """
        else:
            return """
            ## General Study Advice:
            - Create a dedicated study space free from distractions
            - Maintain a consistent study schedule
            - Take regular breaks to avoid burnout
            - Get enough sleep and exercise regularly
            - Don't hesitate to ask for help when needed
            """

    def generate_quiz(self, topic=None, subject=None, difficulty='intermediate', question_count=10, source_type='topic', source_id=None, question_types=None):
        """Generate a quiz based on topic, conversation history, or uploaded documents"""

        if not self.api_available or not self.model:
            return {
                'success': False,
                'error': 'Gemini API not configured. Please set GEMINI_API_KEY in .env file.',
                'fallback_quiz': self._get_fallback_quiz(topic, subject, difficulty, question_count, question_types)
            }

        # Build context based on source type
        context = self._build_quiz_context(source_type, source_id, topic, subject)

        # Handle question types
        if not question_types or len(question_types) == 0:
            question_types = ['multiple_choice', 'true_false', 'short_answer', 'fill_blank']
        question_types_str = ', '.join(question_types)

        # Add randomness to prevent identical questions
        import random
        random_seed = random.randint(1, 10000)

        prompt = f"""
        You are an expert educator creating a quiz for students. Generate {question_count} unique quiz questions.

        Context: {context}
        Topic: {topic or 'General Knowledge'}
        Subject: {subject or 'Various'}
        Difficulty: {difficulty}
        Allowed question types: {question_types_str}

        Generate {question_count} questions using ONLY these types: {question_types_str}

        IMPORTANT:
        - Create varied, unique questions
        - For multiple choice: provide 4 options, correct_answer must match exactly one option
        - For other types: provide the correct answer text

        Respond ONLY with a valid JSON array. Format:
        [
          {{
            "question_type": "multiple_choice",
            "question_text": "Question here?",
            "options": ["A) opt1", "B) opt2", "C) opt3", "D) opt4"],
            "correct_answer": "B) opt2",
            "explanation": "Brief explanation."
          }},
          {{
            "question_type": "short_answer",
            "question_text": "Short answer question?",
            "options": null,
            "correct_answer": "Correct answer here",
            "explanation": "Brief explanation."
          }}
        ]
        """

        try:
            response = self.model.generate_content(prompt)
            response_text = response.text.strip()

            # Clean the response - remove any markdown or extra text
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]

            response_text = response_text.strip()

            # Try to parse JSON with better error handling
            try:
                quiz_data = json.loads(response_text)
            except json.JSONDecodeError as parse_error:
                print(f"JSON parsing failed: {parse_error}")
                print(f"Raw response (first 2000 chars): {response_text[:2000]}...")
                print(f"Raw response (last 2000 chars): {response_text[-2000:]}...")

                # Try to extract JSON from a larger response if it's embedded
                json_start = response_text.find('[')
                json_end = response_text.rfind(']') + 1

                if json_start != -1 and json_end > json_start:
                    try:
                        potential_json = response_text[json_start:json_end]
                        quiz_data = json.loads(potential_json)
                        print("Successfully extracted JSON from larger response")
                    except json.JSONDecodeError:
                        raise ValueError(f"Failed to extract valid JSON from response: {parse_error}")
                else:
                    raise ValueError(f"Failed to parse quiz response: {parse_error}")

            # Validate the structure
            if not isinstance(quiz_data, list):
                raise ValueError("Response is not a JSON array")

            # Ensure we have the right number of questions
            if len(quiz_data) > question_count:
                quiz_data = quiz_data[:question_count]

            return {
                'success': True,
                'questions': quiz_data,
                'topic': topic,
                'difficulty': difficulty
            }
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {e}")
            print(f"Raw response: {response.text[:500]}...")
            return {
                'success': False,
                'error': f'Failed to parse quiz response: {str(e)}',
                'fallback_quiz': self._get_fallback_quiz(topic, subject, difficulty, question_count, question_types)
            }
        except Exception as e:
            print(f"Quiz generation error: {e}")
            return {
                'success': False,
                'error': f'Failed to generate quiz: {str(e)}',
                'fallback_quiz': self._get_fallback_quiz(topic, subject, difficulty, question_count, question_types)
            }

    def _build_quiz_context(self, source_type, source_id, topic, subject):
        """Build context for quiz generation based on source type"""
        context = ""

        try:
            if source_type == 'chat':
                # Generate from recent chat conversations
                from app.models import AIChat
                recent_chats = AIChat.query.filter_by(user_id=source_id).order_by(
                    AIChat.created_at.desc()
                ).limit(20).all()

                context = "Based on recent conversations:\n"
                for chat in recent_chats:
                    context += f"User: {chat.user_message[:200]}...\n"
                    context += f"AI: {chat.ai_response[:300]}...\n\n"

            elif source_type == 'document':
                # Generate from uploaded document
                documents = session.get('uploaded_documents', [])
                if documents:
                    for doc in documents:
                        if doc['filename'] == source_id or str(doc.get('id', '')) == str(source_id):
                            context = f"Based on document '{doc['filename']}':\n{doc['content'][:3000]}..."
                            break

            elif source_type == 'topic':
                # Generate from specified topic
                context = f"Create questions about: {topic or 'General study topics'}"
                if subject:
                    context += f" in the subject area of {subject}"

        except Exception as e:
            print(f"Error building quiz context: {e}")

        return context or "General knowledge quiz"

    def evaluate_answer(self, question_type, user_answer, correct_answer):
        """Evaluate a quiz answer and provide feedback"""
        if not self.api_available or not self.model:
            # Simple evaluation for fallback
            return self._simple_answer_evaluation(question_type, user_answer, correct_answer)

        prompt = f"""
        Evaluate this quiz answer carefully:

        Question Type: {question_type}
        User's Answer: "{user_answer}"
        Correct Answer: "{correct_answer}"

        For multiple choice: Check if the user's answer letter matches the correct answer letter.
        For true/false: Check if the boolean values match.
        For short answer: Check if the meaning is essentially correct (allow minor differences in wording).

        Provide evaluation in JSON format:
        {{
          "is_correct": true/false,
          "feedback": "Brief feedback (max 50 words) explaining the evaluation",
          "points_awarded": 0-1 (1.0 for fully correct, 0 for incorrect)
        }}

        Be fair but accurate in your evaluation.
        """

        try:
            response = self.model.generate_content(prompt)
            evaluation = json.loads(response.text.strip())
            return evaluation
        except Exception as e:
            print(f"Answer evaluation error: {e}")
            return self._simple_answer_evaluation(question_type, user_answer, correct_answer)

    def _simple_answer_evaluation(self, question_type, user_answer, correct_answer):
        """Simple answer evaluation for fallback"""
        user_clean = user_answer.strip().lower() if user_answer else ""
        correct_clean = correct_answer.strip().lower() if correct_answer else ""

        if question_type == 'true_false':
            # Handle true/false answers
            user_bool = user_clean in ['true', 't', 'yes', 'y', '1']
            correct_bool = correct_clean in ['true', 't', 'yes', 'y', '1']
            is_correct = user_bool == correct_bool
        elif question_type == 'multiple_choice':
            # Extract the letter (A, B, C, D) from both answers
            user_letter = user_clean[:1] if user_clean else ""
            correct_letter = correct_clean[:1] if correct_clean else ""
            is_correct = user_letter.lower() == correct_letter.lower()
        else:
            # For short answer and fill_blank, more flexible matching
            # Remove common punctuation and extra spaces
            import re
            user_normalized = re.sub(r'[^\w\s]', '', user_clean).strip()
            correct_normalized = re.sub(r'[^\w\s]', '', correct_clean).strip()

            # Exact match
            if user_normalized == correct_normalized:
                is_correct = True
            # Contains the key answer
            elif len(correct_normalized.split()) <= 3 and all(word in user_normalized for word in correct_normalized.split()):
                is_correct = True
            else:
                is_correct = False

        return {
            'is_correct': is_correct,
            'feedback': 'Correct!' if is_correct else f'Incorrect. The correct answer is: {correct_answer}',
            'points_awarded': 1 if is_correct else 0
        }

    def _get_fallback_quiz(self, topic, subject, difficulty, question_count, question_types=None):
        """Fallback quiz when AI is not available"""
        if not question_types or len(question_types) == 0:
            question_types = ['multiple_choice', 'true_false', 'short_answer', 'fill_blank']

        fallback_questions = []
        topic_lower = topic.lower() if topic else ""

        # Generate questions based on topic
        if 'deep learning' in topic_lower or 'machine learning' in topic_lower or 'ai' in topic_lower:
            # Deep Learning questions
            if 'multiple_choice' in question_types:
                fallback_questions.extend([
                    {
                        'question_type': 'multiple_choice',
                        'question_text': 'What is the primary purpose of the activation function in a neural network?',
                        'options': ['A) To initialize weights', 'B) To introduce non-linearity', 'C) To normalize inputs', 'D) To reduce dimensionality'],
                        'correct_answer': 'B) To introduce non-linearity',
                        'explanation': 'Activation functions introduce non-linearity into the network, allowing it to learn complex patterns.'
                    },
                    {
                        'question_type': 'multiple_choice',
                        'question_text': 'Which optimization algorithm is commonly used in deep learning?',
                        'options': ['A) Linear Regression', 'B) Adam', 'C) K-Means', 'D) Decision Trees'],
                        'correct_answer': 'B) Adam',
                        'explanation': 'Adam (Adaptive Moment Estimation) is a popular optimization algorithm for training neural networks.'
                    },
                    {
                        'question_type': 'multiple_choice',
                        'question_text': 'What does CNN stand for in deep learning?',
                        'options': ['A) Central Neural Network', 'B) Convolutional Neural Network', 'C) Computer Neural Network', 'D) Complex Neural Network'],
                        'correct_answer': 'B) Convolutional Neural Network',
                        'explanation': 'CNN stands for Convolutional Neural Network, commonly used for image processing tasks.'
                    }
                ])

            if 'true_false' in question_types:
                fallback_questions.extend([
                    {
                        'question_type': 'true_false',
                        'question_text': 'Backpropagation is used to update the weights in a neural network.',
                        'options': None,
                        'correct_answer': 'true',
                        'explanation': 'Backpropagation calculates the gradient of the loss function and updates network weights.'
                    },
                    {
                        'question_type': 'true_false',
                        'question_text': 'Overfitting occurs when a model performs well on training data but poorly on new data.',
                        'options': None,
                        'correct_answer': 'true',
                        'explanation': 'Overfitting happens when a model learns the training data too well, including noise and outliers.'
                    }
                ])

            if 'short_answer' in question_types:
                fallback_questions.extend([
                    {
                        'question_type': 'short_answer',
                        'question_text': 'What is the main difference between supervised and unsupervised learning?',
                        'options': None,
                        'correct_answer': 'Supervised learning uses labeled data, unsupervised learning uses unlabeled data',
                        'explanation': 'Supervised learning requires labeled training data, while unsupervised learning finds patterns in unlabeled data.'
                    },
                    {
                        'question_type': 'short_answer',
                        'question_text': 'Name one technique used to prevent overfitting in deep learning.',
                        'options': None,
                        'correct_answer': 'Dropout',
                        'explanation': 'Dropout randomly deactivates neurons during training to prevent overfitting.'
                    }
                ])


        elif 'math' in topic_lower:
            # Math questions
            if 'multiple_choice' in question_types:
                fallback_questions.extend([
                    {
                        'question_type': 'multiple_choice',
                        'question_text': 'What is the derivative of sin(x)?',
                        'options': ['A) cos(x)', 'B) -cos(x)', 'C) tan(x)', 'D) -sin(x)'],
                        'correct_answer': 'A) cos(x)',
                        'explanation': 'The derivative of sin(x) is cos(x).'
                    },
                    {
                        'question_type': 'multiple_choice',
                        'question_text': 'What is the value of π (pi) to 3 decimal places?',
                        'options': ['A) 3.142', 'B) 3.141', 'C) 3.140', 'D) 3.143'],
                        'correct_answer': 'B) 3.141',
                        'explanation': 'π is approximately 3.14159, so to 3 decimal places it is 3.142.'
                    }
                ])

            if 'short_answer' in question_types:
                fallback_questions.extend([
                    {
                        'question_type': 'short_answer',
                        'question_text': 'Solve for x: 2x + 3 = 7',
                        'options': None,
                        'correct_answer': 'x = 2',
                        'explanation': 'Subtract 3 from both sides: 2x = 4, then divide by 2: x = 2.'
                    }
                ])

        else:
            # General knowledge questions
            if 'multiple_choice' in question_types:
                fallback_questions.extend([
                    {
                        'question_type': 'multiple_choice',
                        'question_text': 'What is the capital of France?',
                        'options': ['A) London', 'B) Paris', 'C) Berlin', 'D) Madrid'],
                        'correct_answer': 'B) Paris',
                        'explanation': 'Paris is the capital and largest city of France.'
                    },
                    {
                        'question_type': 'multiple_choice',
                        'question_text': 'Which planet is known as the Red Planet?',
                        'options': ['A) Venus', 'B) Mars', 'C) Jupiter', 'D) Saturn'],
                        'correct_answer': 'B) Mars',
                        'explanation': 'Mars is called the Red Planet due to its reddish appearance caused by iron oxide on its surface.'
                    },
                    {
                        'question_type': 'multiple_choice',
                        'question_text': 'What is the largest ocean on Earth?',
                        'options': ['A) Atlantic Ocean', 'B) Indian Ocean', 'C) Arctic Ocean', 'D) Pacific Ocean'],
                        'correct_answer': 'D) Pacific Ocean',
                        'explanation': 'The Pacific Ocean is the largest ocean on Earth, covering about 46% of the world\'s water surface.'
                    }
                ])

            if 'true_false' in question_types:
                fallback_questions.extend([
                    {
                        'question_type': 'true_false',
                        'question_text': 'The Earth revolves around the Sun.',
                        'options': None,
                        'correct_answer': 'true',
                        'explanation': 'The Earth orbits the Sun in our solar system.'
                    },
                    {
                        'question_type': 'true_false',
                        'question_text': 'Water boils at 100 degrees Celsius at sea level.',
                        'options': None,
                        'correct_answer': 'true',
                        'explanation': 'At standard atmospheric pressure (sea level), water boils at 100°C.'
                    }
                ])

            if 'short_answer' in question_types:
                fallback_questions.extend([
                    {
                        'question_type': 'short_answer',
                        'question_text': 'What is the chemical symbol for gold?',
                        'options': None,
                        'correct_answer': 'Au',
                        'explanation': 'Au comes from the Latin word "aurum" meaning gold.'
                    },
                    {
                        'question_type': 'short_answer',
                        'question_text': 'What is the largest planet in our solar system?',
                        'options': None,
                        'correct_answer': 'Jupiter',
                        'explanation': 'Jupiter is the largest planet in our solar system by mass and volume.'
                    }
                ])


        # If we don't have enough questions, repeat some with variations
        while len(fallback_questions) < question_count:
            for q in fallback_questions[:]:
                if len(fallback_questions) >= question_count:
                    break
                # Create a variation by modifying slightly
                if q['question_type'] == 'multiple_choice':
                    fallback_questions.append(q.copy())
                elif len(fallback_questions) < question_count:
                    fallback_questions.append(q.copy())

        # Shuffle and limit to requested count
        import random
        random.shuffle(fallback_questions)

        return {
            'title': f'{topic or "General"} Quiz',
            'questions': fallback_questions[:question_count],
            'difficulty': difficulty
        }

# Initialize the AI tutor
gemini_tutor = GeminiTutor()

@ai_tutor_bp.route('/')
@login_required
def index():
    """AI Tutor main chat interface"""
    # Get recent conversations
    recent_chats = AIChat.query.filter_by(user_id=current_user.id).order_by(
        AIChat.created_at.desc()
    ).limit(50).all()

    # Group chats by date
    chats_by_date = {}
    for chat in recent_chats:
        date_key = chat.created_at.strftime('%Y-%m-%d')
        if date_key not in chats_by_date:
            chats_by_date[date_key] = []
        chats_by_date[date_key].append(chat)

    return render_template('ai_tutor/index.html',
                           title='AI Tutor',
                           chats_by_date=chats_by_date)

@ai_tutor_bp.route('/chat', methods=['POST'])
@login_required
def chat():
    """Handle chat messages"""
    user_message = request.form.get('message', '').strip()
    message_type = request.form.get('message_type', 'general')
    context = request.form.get('context', '')
    language = request.form.get('language', 'en')

    if not user_message:
        return jsonify({'error': 'Message cannot be empty'}), 400

    # Get AI response
    try:
        result = gemini_tutor.get_response(user_message, message_type, context, language)

        # Handle the response format
        if result['success']:
            ai_response = result['content']
        else:
            ai_response = result['fallback_content']
            print(f"DEBUG: Using fallback response: {result['error']}")

    except Exception as e:
        ai_response = f"I apologize, but I encountered an error processing your request. Please try rephrasing your question. Error: {str(e)}"
        print(f"DEBUG: Unexpected error in chat route: {e}")

    # Save conversation to database
    chat_entry = AIChat(
        user_id=current_user.id,
        user_message=user_message,
        ai_response=ai_response,
        message_type=message_type,
        context=context
    )

    try:
        db.session.add(chat_entry)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': 'Failed to save conversation'}), 500

    return jsonify({
        'response': ai_response,
        'chat_id': chat_entry.id,
        'timestamp': chat_entry.created_at.strftime('%H:%M')
    })

@ai_tutor_bp.route('/clear-history', methods=['POST'])
@login_required
def clear_history():
    """Clear user's chat history"""
    try:
        AIChat.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        flash('Chat history cleared successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash('Failed to clear chat history.', 'error')

    return redirect(url_for('ai_tutor.index'))

@ai_tutor_bp.route('/study-tip')
@login_required
def study_tip():
    """Get a random study tip"""
    tip = random.choice(gemini_tutor.study_tips)
    return jsonify({'tip': tip})

@ai_tutor_bp.route('/suggest-topics')
@login_required
def suggest_topics():
    """Suggest study topics based on user's goals and tasks"""
    # Get user's active goals and recent tasks
    from app.models import Goal, Task

    active_goals = Goal.query.filter_by(
        user_id=current_user.id,
        achieved=False
    ).limit(5).all()

    recent_tasks = Task.query.filter_by(
        user_id=current_user.id
    ).order_by(Task.created_at.desc()).limit(10).all()

    topics = []

    # Extract topics from goals
    for goal in active_goals:
        if goal.category:
            topics.append(f"{goal.category.title()} studies")
        topics.append(f"Goal: {goal.title}")

    # Extract topics from tasks
    for task in recent_tasks:
        if task.category:
            topics.append(f"{task.category.title()} topics")
        if task.title:
            topics.append(f"Task: {task.title}")

    # Remove duplicates and limit
    unique_topics = list(set(topics))[:10]

    if not unique_topics:
        unique_topics = [
            "Time management techniques",
            "Study skill improvement",
            "Goal setting strategies",
            "Memory and retention methods",
            "Exam preparation tips"
        ]

    return jsonify({'topics': unique_topics})

@ai_tutor_bp.route('/upload-document', methods=['POST'])
@login_required
def upload_document():
    """Handle document upload and analysis"""
    logging.info(f"Upload route called for user {current_user.id}")
    try:
        logging.info(f"Starting upload processing for user {current_user.id}")
        if 'file' not in request.files:
            logging.warning("No file in request.files")
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        logging.info(f"File received: {file.filename}, size: {getattr(file, 'content_length', 'unknown')}")
        if file.filename == '':
            logging.warning("Empty filename")
            return jsonify({'error': 'No file selected', 'debug': 'Empty filename provided'}), 400

        # Validate file type and size
        allowed_extensions = {'pdf', 'docx', 'txt'}
        max_file_size = 50 * 1024 * 1024  # 50MB

        if not file or not allowed_filename(file.filename, allowed_extensions):
            logging.warning(f"Invalid file type: {file.filename}")
            return jsonify({'error': 'Invalid file type. Only PDF, DOCX, and TXT files are allowed.', 'debug': f'File: {file.filename}'}), 400

        if file.content_length and file.content_length > max_file_size:
            logging.warning(f"File too large: {file.content_length}")
            return jsonify({'error': 'File too large. Maximum size is 50MB.', 'debug': f'Size: {file.content_length}'}), 400

        # Save file temporarily for processing
        import tempfile

        # Use system temp directory for better reliability
        temp_dir = os.path.join(tempfile.gettempdir(), 'ai_tutor_temp')
        os.makedirs(temp_dir, exist_ok=True)

        filename = secure_filename(file.filename)
        file_path = os.path.join(temp_dir, filename)
        print(f"DEBUG: Saving file to {file_path}")
        file.save(file_path)

        # Extract text from document with timeout
        print(f"DEBUG: Extracting text from {filename}")
        document_text = extract_text_from_file(file_path, filename, timeout_seconds=28)

        # Clean up temp file
        os.remove(file_path)
        print(f"DEBUG: Cleaned up temp file {file_path}")

        if not document_text:
            file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'unknown'
            if file_ext == 'doc':
                error_msg = 'DOC files are not supported. Please convert your Word document to DOCX format (Save As → Word Document) and try again.'
            elif file_ext == 'docx':
                error_msg = 'Could not extract text from the DOCX file. The file may be corrupted or password-protected.'
            elif file_ext == 'pdf':
                error_msg = 'Could not extract text from the PDF file. The file may be scanned/image-based, corrupted, or password-protected.'
            else:
                error_msg = 'Could not extract text from the document. The file may be corrupted, password-protected, or contain only images/scanned content.'
            print(f"DEBUG: No text extracted from {filename} (type: {file_ext})")
            return jsonify({'error': error_msg}), 400

        print(f"DEBUG: Extracted {len(document_text)} characters from {filename}")

        # Store document in session for context (simplified approach)
        # In production, you'd store this in a database
        if 'uploaded_documents' not in session:
            session['uploaded_documents'] = []
            print("DEBUG: Initialized uploaded_documents in session")

        # Limit to 4 documents max
        if len(session['uploaded_documents']) >= 4:
            print(f"DEBUG: Too many documents: {len(session['uploaded_documents'])}")
            return jsonify({'error': 'Maximum 4 documents allowed. Please remove some files before uploading more.'}), 400

        doc_info = {
            'filename': filename,
            'content': document_text[:2000],  # Store first 2000 chars for context
            'uploaded_at': datetime.utcnow().isoformat()
        }

        session['uploaded_documents'].append(doc_info)
        session.modified = True
        print(f"DEBUG: Stored document in session. Total documents: {len(session['uploaded_documents'])}")
        print(f"DEBUG: Session contents: {session.get('uploaded_documents')}")

        print(f"DEBUG: Upload successful for {filename}, returning response")
        return jsonify({
            'success': True,
            'message': f'Successfully uploaded and analyzed {filename}',
            'filename': filename,
            'text_length': len(document_text)
        })

    except Exception as e:
        print(f"DEBUG: Upload failed with exception: {str(e)}")
        import traceback
        print(f"DEBUG: Traceback: {traceback.format_exc()}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

def allowed_filename(filename, allowed_extensions):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

@ai_tutor_bp.route('/remove-document/<filename>', methods=['POST'])
@login_required
def remove_document(filename):
    """Remove a document from session"""
    try:
        if 'uploaded_documents' in session:
            session['uploaded_documents'] = [
                doc for doc in session['uploaded_documents']
                if doc['filename'] != filename
            ]
            session.modified = True
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@ai_tutor_bp.route('/get-uploaded-documents')
@login_required
def get_uploaded_documents():
    """Get list of uploaded documents"""
    try:
        documents = session.get('uploaded_documents', [])
        print(f"DEBUG: Get uploaded documents for user {current_user.id}: {documents}")
        return jsonify({'documents': documents})
    except Exception as e:
        print(f"DEBUG: Error getting uploaded documents: {str(e)}")
        return jsonify({'error': str(e)}), 500

@ai_tutor_bp.route('/test-upload')
@login_required
def test_upload():
    """Simple test route to check if upload endpoint is accessible"""
    return jsonify({
        'status': 'ok',
        'message': 'Upload endpoint is accessible',
        'user': current_user.id,
        'timestamp': datetime.utcnow().isoformat()
    })

@ai_tutor_bp.route('/generate-quiz', methods=['POST'])
@login_required
def generate_quiz():
    """Generate a quiz from various sources"""
    try:
        data = request.get_json() or {}
        topic = data.get('topic', '')
        subject = data.get('subject', '')
        difficulty = data.get('difficulty', 'intermediate')
        question_count = min(int(data.get('question_count', 10)), 20)  # Max 20 questions
        source_type = data.get('source_type', 'topic')
        source_id = data.get('source_id')
        question_types = data.get('question_types', ['multiple_choice', 'true_false', 'short_answer'])

        # Ensure question_types is a list
        if isinstance(question_types, str):
            question_types = [question_types]
        elif not question_types:
            question_types = ['multiple_choice', 'true_false', 'short_answer']

        # Generate quiz using AI
        result = gemini_tutor.generate_quiz(topic, subject, difficulty, question_count, source_type, source_id, question_types)

        if not result['success']:
            return jsonify({'error': result['error']}), 500

        # Save quiz to database
        quiz = Quiz(
            user_id=current_user.id,
            title=f"{topic or 'General'} Quiz" if not subject else f"{subject}: {topic or 'Quiz'}",
            topic=topic,
            subject=subject,
            difficulty=difficulty,
            question_count=question_count,
            generated_from=source_type,
            source_id=source_id
        )

        db.session.add(quiz)
        db.session.flush()  # Get quiz ID

        # Save questions with validation
        max_score = 0
        valid_questions_count = 0

        print(f"Processing {len(result['questions'])} questions for quiz generation")

        for i, q_data in enumerate(result['questions']):
            print(f"Processing question {i}: {q_data.get('question_type', 'unknown type')}")

            # Validate required fields with better error handling
            question_text = q_data.get('question_text', '').strip()
            question_type = q_data.get('question_type', '').strip()
            correct_answer = q_data.get('correct_answer')

            if not question_text:
                print(f"Skipping question {i}: empty question_text")
                continue

            if not question_type:
                print(f"Skipping question {i}: empty question_type")
                continue

            # Ensure question_type is valid
            valid_types = ['multiple_choice', 'true_false', 'short_answer', 'fill_blank']
            if question_type not in valid_types:
                print(f"Skipping question {i}: invalid question_type '{question_type}'")
                continue

            # For multiple choice, ensure options exist
            if question_type == 'multiple_choice':
                if not q_data.get('options') or not isinstance(q_data['options'], list) or len(q_data['options']) < 2:
                    print(f"Skipping multiple choice question {i}: missing or invalid options")
                    continue
                # Ensure correct_answer matches one of the options
                if correct_answer not in q_data['options']:
                    print(f"Skipping multiple choice question {i}: correct_answer '{correct_answer}' not in options")
                    continue

            # Validate correct_answer based on question type
            if correct_answer is None or correct_answer == '':
                print(f"Skipping question {i}: empty correct_answer")
                continue

            # Convert correct_answer to string and strip whitespace
            correct_answer = str(correct_answer).strip()
            if not correct_answer:
                print(f"Skipping question {i}: correct_answer is only whitespace")
                continue

            question = QuizQuestion(
                quiz_id=quiz.id,
                question_text=question_text,
                question_type=question_type,
                options=q_data.get('options'),
                correct_answer=correct_answer,
                explanation=q_data.get('explanation', '').strip() if q_data.get('explanation') else '',
                points=q_data.get('points', 1),
                order=i + 1
            )
            db.session.add(question)
            max_score += question.points
            valid_questions_count += 1
            print(f"Successfully added question {i}: {q_data['question_type']}")

        print(f"Total valid questions saved: {valid_questions_count}, max_score: {max_score}")

        # Check if we have at least one valid question
        if valid_questions_count == 0:
            db.session.rollback()
            return jsonify({'error': 'Failed to generate any valid quiz questions. All questions were filtered out due to validation issues. Please try again.'}), 500

        # If we have fewer than requested questions, log it but continue
        if valid_questions_count < question_count:
            print(f"Warning: Only {valid_questions_count} out of {question_count} requested questions were valid and saved.")

        quiz.max_score = max_score
        db.session.commit()

        return jsonify({
            'success': True,
            'quiz_id': quiz.id,
            'title': quiz.title,
            'question_count': len(result['questions']),
            'difficulty': difficulty
        })

    except Exception as e:
        db.session.rollback()
        print(f"Quiz generation error: {e}")
        return jsonify({'error': f'Failed to generate quiz: {str(e)}'}), 500

@ai_tutor_bp.route('/quizzes')
@login_required
def get_quizzes():
    """Get user's quizzes"""
    try:
        quizzes = Quiz.query.filter_by(user_id=current_user.id).order_by(Quiz.created_at.desc()).all()

        quiz_list = []
        for quiz in quizzes:
            quiz_list.append({
                'id': quiz.id,
                'title': quiz.title,
                'topic': quiz.topic,
                'difficulty': quiz.difficulty,
                'question_count': quiz.question_count,
                'created_at': quiz.created_at.strftime('%Y-%m-%d %H:%M'),
                'attempts_count': len(quiz.attempts)
            })

        return jsonify({'quizzes': quiz_list})

    except Exception as e:
        print(f"Error getting quizzes: {e}")
        return jsonify({'error': str(e)}), 500

@ai_tutor_bp.route('/quiz/<int:quiz_id>')
@login_required
def get_quiz(quiz_id):
    """Get quiz details and questions"""
    try:
        quiz = Quiz.query.filter_by(id=quiz_id, user_id=current_user.id).first()
        if not quiz:
            return jsonify({'error': 'Quiz not found'}), 404

        questions = []
        for question in quiz.questions:
            questions.append({
                'id': question.id,
                'question_text': question.question_text,
                'question_type': question.question_type,
                'options': question.options,
                'correct_answer': question.correct_answer,
                'explanation': question.explanation,
                'points': question.points
            })

        return jsonify({
            'id': quiz.id,
            'title': quiz.title,
            'topic': quiz.topic,
            'subject': quiz.subject,
            'difficulty': quiz.difficulty,
            'time_limit': quiz.time_limit,
            'questions': questions,
            'max_score': quiz.max_score
        })

    except Exception as e:
        print(f"Error getting quiz: {e}")
        return jsonify({'error': str(e)}), 500

@ai_tutor_bp.route('/quiz/<int:quiz_id>/attempt', methods=['POST'])
@login_required
def start_quiz_attempt(quiz_id):
    """Start a quiz attempt"""
    try:
        quiz = Quiz.query.filter_by(id=quiz_id, user_id=current_user.id).first()
        if not quiz:
            return jsonify({'error': 'Quiz not found'}), 404

        # Calculate max_score from questions (always calculate fresh)
        total_points = sum(question.points for question in quiz.questions)

        # Update quiz max_score if it's different
        if quiz.max_score != total_points:
            quiz.max_score = total_points
            db.session.commit()

        # Create new attempt
        attempt = QuizAttempt(
            user_id=current_user.id,
            quiz_id=quiz.id,
            max_score=total_points
        )

        db.session.add(attempt)
        db.session.commit()

        return jsonify({
            'attempt_id': attempt.id,
            'quiz_title': quiz.title,
            'time_limit': quiz.time_limit,
            'question_count': quiz.question_count
        })

    except AttributeError as e:
        # Handle case where quiz object doesn't have expected attributes
        db.session.rollback()
        print(f"Attribute error starting quiz attempt: {e}")
        return jsonify({'error': 'Quiz data is corrupted. Please try creating a new quiz.'}), 500
    except Exception as e:
        db.session.rollback()
        print(f"Error starting quiz attempt: {e}")
        return jsonify({'error': str(e)}), 500

@ai_tutor_bp.route('/quiz/attempt/<int:attempt_id>/answer', methods=['POST'])
@login_required
def submit_answer(attempt_id):
    """Submit an answer for a quiz question"""
    try:
        data = request.get_json()
        question_id = data.get('question_id')
        user_answer = data.get('answer', '').strip()

        # Get attempt and verify ownership
        attempt = QuizAttempt.query.filter_by(id=attempt_id, user_id=current_user.id).first()
        if not attempt:
            return jsonify({'error': 'Attempt not found'}), 404

        # Get question
        question = QuizQuestion.query.filter_by(id=question_id, quiz_id=attempt.quiz_id).first()
        if not question:
            return jsonify({'error': 'Question not found'}), 404

        # Evaluate answer
        evaluation = gemini_tutor.evaluate_answer(question.question_type, user_answer, question.correct_answer)

        # Save answer
        answer = QuizAnswer(
            attempt_id=attempt.id,
            question_id=question.id,
            user_answer=user_answer,
            is_correct=evaluation['is_correct'],
            points_earned=int(evaluation['points_awarded'] * question.points)
        )

        db.session.add(answer)

        # Update attempt score
        attempt.score += answer.points_earned
        attempt.percentage = (attempt.score / attempt.max_score) * 100 if attempt.max_score > 0 else 0

        db.session.commit()

        return jsonify({
            'is_correct': evaluation['is_correct'],
            'feedback': evaluation['feedback'],
            'points_earned': answer.points_earned,
            'current_score': attempt.score
        })

    except Exception as e:
        db.session.rollback()
        print(f"Error submitting answer: {e}")
        return jsonify({'error': str(e)}), 500

@ai_tutor_bp.route('/quiz/attempt/<int:attempt_id>/complete', methods=['POST'])
@login_required
def complete_quiz_attempt(attempt_id):
    """Complete a quiz attempt"""
    try:
        attempt = QuizAttempt.query.filter_by(id=attempt_id, user_id=current_user.id).first()
        if not attempt:
            return jsonify({'error': 'Attempt not found'}), 404

        # Calculate time taken
        time_taken = (datetime.utcnow() - attempt.started_at).total_seconds()

        # Update attempt
        attempt.completed_at = datetime.utcnow()
        attempt.time_taken = int(time_taken)
        attempt.completed = True

        db.session.commit()

        return jsonify({
            'score': attempt.score,
            'max_score': attempt.max_score,
            'percentage': round(attempt.percentage, 1),
            'time_taken': attempt.time_taken
        })

    except Exception as e:
        db.session.rollback()
        print(f"Error completing quiz attempt: {e}")
        return jsonify({'error': str(e)}), 500

def extract_text_from_file(file_path, filename, timeout_seconds=25):
    """Extract text from various document formats with timeout"""
    import time
    start_time = time.time()

    def check_timeout():
        if time.time() - start_time > timeout_seconds:
            raise TimeoutError(f"Text extraction timed out after {timeout_seconds} seconds")

    try:
        file_extension = filename.rsplit('.', 1)[1].lower()

        if file_extension == 'txt':
            try:
                check_timeout()
                # First try to detect if it's actually a text file
                with open(file_path, 'rb') as f:
                    sample = f.read(1024)  # Read first 1KB
                    # More lenient check - allow up to 20% null bytes or control characters
                    null_count = sample.count(b'\x00')
                    control_count = sum(1 for b in sample if b < 32 and b not in [9, 10, 13])  # Tab, LF, CR allowed
                    if null_count > len(sample) * 0.2 and control_count > len(sample) * 0.3:
                        print(f"DEBUG: File {filename} appears to be binary (null: {null_count}, control: {control_count}), skipping")
                        return None

                check_timeout()
                # Try different encodings with more attempts
                encodings = ['utf-8', 'utf-16', 'utf-32', 'latin-1', 'cp1252', 'iso-8859-1', 'windows-1252']
                text = None
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                            text = f.read()
                            # Clean up replacement characters and check if we have meaningful content
                            cleaned_text = text.replace('\ufffd', '').strip()
                            if len(cleaned_text) > 10:  # At least 10 characters of actual content
                                print(f"DEBUG: Successfully read TXT file with {encoding} encoding, length: {len(cleaned_text)}")
                                return cleaned_text
                    except (UnicodeDecodeError, UnicodeError, LookupError):
                        continue

                # If no encoding worked, try reading as binary and decode with errors='replace'
                if not text:
                    try:
                        with open(file_path, 'rb') as f:
                            binary_content = f.read()
                            # Try to decode as utf-8 with replace
                            text = binary_content.decode('utf-8', errors='replace')
                            cleaned_text = text.replace('\ufffd', '').strip()
                            if len(cleaned_text) > 10:
                                print(f"DEBUG: Read TXT file as binary fallback, length: {len(cleaned_text)}")
                                return cleaned_text
                    except Exception:
                        pass

                return None
            except Exception as e:
                print(f"TXT extraction error for {filename}: {str(e)}")
                return None

        elif file_extension == 'pdf':
            try:
                check_timeout()
                from pypdf import PdfReader
                with open(file_path, 'rb') as f:
                    pdf_reader = PdfReader(f)
                    text = ''
                    max_pages = min(len(pdf_reader.pages), 100)  # Limit to 100 pages to prevent timeout
                    for i in range(max_pages):
                        check_timeout()
                        page_text = pdf_reader.pages[i].extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text += page_text + '\n'
                    return text.strip() if text.strip() else None
            except TimeoutError:
                print(f"PDF extraction timed out for {filename}")
                return None
            except Exception as e:
                print(f"PDF extraction error for {filename}: {str(e)}")
                return None

        elif file_extension == 'docx':
            try:
                check_timeout()
                from docx import Document
                doc = Document(file_path)
                text = ''
                check_timeout()
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Only add non-empty paragraphs
                        text += paragraph.text + '\n'

                check_timeout()
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + '\n'

                return text.strip() if text.strip() else None
            except Exception as e:
                print(f"DOCX extraction error for {filename}: {str(e)}")
                return None

        elif file_extension == 'doc':
            # DOC files not supported - user should convert to DOCX
            return None

        else:
            return None

    except TimeoutError:
        print(f"Text extraction timed out for {filename}")
        return None
    except Exception as e:
        print(f"General extraction error for {filename}: {str(e)}")
        return None